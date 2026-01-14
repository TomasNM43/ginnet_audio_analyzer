import sys
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QLabel, QLineEdit, QPushButton, QFileDialog, QMessageBox, QRadioButton, QButtonGroup, QListWidget, QHBoxLayout, QComboBox
import mysql.connector as conn  # Importamos el conector MySQL
import os
import shutil
import librosa
import librosa.display
import matplotlib.pyplot as plt
import numpy as np
import cv2
from ultralytics import YOLO
from docx import Document
from docx.shared import Inches
from PIL import Image
import datetime
import speech_recognition as sr
from pydub import AudioSegment
import tempfile
import soundfile as sf

# Función para guardar el espectrograma
def save_spectrogram(y, sr, start, end, start_time, end_time, output_path):
    plt.figure(figsize=(10, 4))
    D = librosa.amplitude_to_db(np.abs(librosa.stft(y[start:end])), ref=np.max)
    librosa.display.specshow(D, sr=sr, x_axis='time', y_axis='log', cmap='gray')  # Escala de grises
    plt.colorbar(format='%+2.0f dB')
    plt.title(f'Spectrogram: {int(start_time // 60)}:{int(start_time % 60):02d} - {int(end_time // 60)}:{int(end_time % 60):02d} min')
    plt.tight_layout()
    plt.savefig(output_path)
    plt.close()

# Función para generar los espectrogramas con prefijo de archivo
def generate_spectrograms_for_file(audio_path, file_prefix, segment_length=3, extra_duration=0.05):
    y, sr = librosa.load(audio_path, sr=None)
    duration = librosa.get_duration(y=y, sr=sr)
    
    output_dir = 'spectrograms'
    os.makedirs(output_dir, exist_ok=True)

    i = 0
    while i < duration:
        start_time = i
        end_time = i + segment_length + extra_duration
        start_sample = int(start_time * sr)
        end_sample = int(end_time * sr)
        
        if end_sample > len(y):
            end_sample = len(y)
            end_time = len(y) / sr  # Ajustar el tiempo final si se llega al final del audio
        
        if end_sample > start_sample:
            output_path = os.path.join(output_dir, f'{file_prefix}_spectrogram_{int(start_time)}_{int(end_time)}.png')
            save_spectrogram(y[start_sample:end_sample], sr, 0, end_sample - start_sample, start_time, end_time, output_path)
        
        i += segment_length  # Avanzar al siguiente segmento

# Función para generar espectrogramas en un rango de tiempo específico
def generate_spectrograms_by_time_range(audio_path, file_prefix, start_time, end_time, segment_length=3, mode="complete", time_jump=3, extra_duration=0.05):
    y, sr = librosa.load(audio_path, sr=None)
    duration = librosa.get_duration(y=y, sr=sr)
    
    # Validar que los tiempos estén dentro del rango del audio
    if start_time < 0:
        start_time = 0
    if end_time > duration:
        end_time = duration
    if start_time >= end_time:
        raise ValueError("El tiempo inicial debe ser menor que el tiempo final")
    
    # Directorio específico para espectrogramas de rango de tiempo
    output_dir = 'spectrograms_time_range'
    os.makedirs(output_dir, exist_ok=True)

    i = start_time
    
    if mode == "complete":
        # Modo: Cobertura completa del rango (todos los fotogramas)
        advance_step = segment_length
    else:  # mode == "combined"
        # Modo: Combinado - primero cobertura completa, luego saltos adicionales
        # Primero hacer cobertura completa
        temp_i = start_time
        while temp_i < end_time:
            segment_start_time = temp_i
            segment_end_time = min(temp_i + segment_length + extra_duration, end_time)
            start_sample = int(segment_start_time * sr)
            end_sample = int(segment_end_time * sr)
            
            if end_sample > len(y):
                end_sample = len(y)
                segment_end_time = len(y) / sr
            
            if end_sample > start_sample:
                output_path = os.path.join(output_dir, f'{file_prefix}_complete_spectrogram_{int(segment_start_time)}_{int(segment_end_time)}.png')
                save_spectrogram(y[start_sample:end_sample], sr, 0, end_sample - start_sample, segment_start_time, segment_end_time, output_path)
            
            temp_i += segment_length
        
        # Luego hacer saltos específicos (evitando duplicados)
        advance_step = time_jump

    while i < end_time:
        segment_start_time = i
        segment_end_time = min(i + segment_length + extra_duration, end_time)
        start_sample = int(segment_start_time * sr)
        end_sample = int(segment_end_time * sr)
        
        if end_sample > len(y):
            end_sample = len(y)
            segment_end_time = len(y) / sr
        
        if end_sample > start_sample:
            if mode == "combined":
                # Evitar duplicados en modo combinado
                if i % segment_length != 0 or i == start_time:
                    prefix = "jump"
                else:
                    i += advance_step
                    continue
            else:
                prefix = "range"
            
            output_path = os.path.join(output_dir, f'{file_prefix}_{prefix}_spectrogram_{int(segment_start_time)}_{int(segment_end_time)}.png')
            save_spectrogram(y[start_sample:end_sample], sr, 0, end_sample - start_sample, segment_start_time, segment_end_time, output_path)
        
        i += advance_step

# Función para generar espectrogramas solo por saltos específicos
def generate_spectrograms_by_jumps(audio_path, file_prefix, time_jump=3, segment_length=3, extra_duration=0.05):
    y, sr = librosa.load(audio_path, sr=None)
    duration = librosa.get_duration(y=y, sr=sr)
    
    # Directorio específico para espectrogramas por saltos
    output_dir = 'spectrograms_jumps'
    os.makedirs(output_dir, exist_ok=True)

    i = 0
    while i < duration:
        segment_start_time = i
        segment_end_time = min(i + segment_length + extra_duration, duration)
        start_sample = int(segment_start_time * sr)
        end_sample = int(segment_end_time * sr)
        
        if end_sample > len(y):
            end_sample = len(y)
            segment_end_time = len(y) / sr
        
        if end_sample > start_sample:
            output_path = os.path.join(output_dir, f'{file_prefix}_jump_spectrogram_{int(segment_start_time)}_{int(segment_end_time)}.png')
            save_spectrogram(y[start_sample:end_sample], sr, 0, end_sample - start_sample, segment_start_time, segment_end_time, output_path)
        
        i += time_jump  # Avanzar según el salto de tiempo configurado

# Función para convertir audio a WAV (necesario para speech_recognition)
def convert_to_wav(audio_path):
    """Convierte cualquier formato de audio a WAV temporal"""
    try:
        # Cargar audio con pydub
        if audio_path.lower().endswith('.wav'):
            return audio_path  # Ya es WAV, no necesita conversión
        
        audio = AudioSegment.from_file(audio_path)
        
        # Crear archivo temporal WAV
        temp_wav = tempfile.NamedTemporaryFile(suffix='.wav', delete=False)
        audio.export(temp_wav.name, format='wav')
        temp_wav.close()
        
        return temp_wav.name
    except Exception as e:
        print(f"Error convirtiendo audio {audio_path}: {e}")
        return None

# Función para transcribir audio completo con múltiples métodos
def transcribe_audio(audio_path, language='es-ES', max_duration=300):
    """Transcribe un archivo de audio completo usando varios métodos de reconocimiento"""
    recognizer = sr.Recognizer()
    
    # Configurar el reconocedor para mejor precisión
    recognizer.energy_threshold = 300
    recognizer.dynamic_energy_threshold = True
    recognizer.pause_threshold = 0.8
    recognizer.operation_timeout = None
    recognizer.phrase_threshold = 0.3
    recognizer.non_speaking_duration = 0.8
    
    # Convertir a WAV si es necesario
    wav_path = convert_to_wav(audio_path)
    if wav_path is None:
        return None, "Error al convertir el archivo de audio"
    
    try:
        # Verificar duración del archivo
        try:
            y, sr_lib = librosa.load(wav_path, sr=None)
            duration = librosa.get_duration(y=y, sr=sr_lib)
            
            if duration > max_duration:
                return f"[Archivo demasiado largo ({duration:.1f}s). Máximo permitido: {max_duration}s. Use la función de segmentación.]", "Archivo muy largo"
        except:
            pass
        
        # Cargar archivo de audio
        with sr.AudioFile(wav_path) as source:
            # Ajustar para ruido ambiente con más tiempo
            recognizer.adjust_for_ambient_noise(source, duration=1)
            # Grabar el audio completo
            audio_data = recognizer.record(source)
        
        # Método 1: Google Speech Recognition con configuración mejorada
        try:
            text = recognizer.recognize_google(
                audio_data, 
                language=language,
                show_all=False
            )
            return text, "Google Speech Recognition"
        except sr.UnknownValueError:
            pass  # Intentar siguiente método
        except sr.RequestError as e:
            print(f"Error Google API: {e}")
            # Continuar con otros métodos
        
        # Método 2: Wit.ai (requiere clave API pero es más flexible)
        try:
            # Nota: Requiere configurar WIT_AI_KEY como variable de entorno
            text = recognizer.recognize_wit(audio_data, key="TU_WIT_AI_KEY")
            return text, "Wit.ai"
        except (sr.UnknownValueError, sr.RequestError):
            pass  # Intentar siguiente método
        
        # Método 3: Reconocimiento por segmentos para archivos largos
        try:
            return transcribe_by_segments(wav_path, recognizer, language)
        except:
            pass
        
        # Método 4: CMU Sphinx (offline, menos preciso pero funciona sin internet)
        try:
            text = recognizer.recognize_sphinx(audio_data, language='es-ES')
            return text, "CMU Sphinx (offline)"
        except (sr.UnknownValueError, sr.RequestError):
            pass
        
        # Si todos los métodos fallan
        return "[No se pudo transcribir el audio con ningún método disponible]", "Error - todos los métodos fallaron"
    
    except Exception as e:
        return f"[Error procesando archivo: {str(e)}]", "Error de procesamiento"
    
    finally:
        # Limpiar archivo temporal si se creó
        if wav_path != audio_path and os.path.exists(wav_path):
            try:
                os.unlink(wav_path)
            except:
                pass

# Función para transcribir por segmentos (archivos largos)
def transcribe_by_segments(wav_path, recognizer, language='es-ES', segment_length=30):
    """Transcribe archivos largos dividiéndolos en segmentos"""
    try:
        # Cargar audio con librosa para segmentar
        y, sr = librosa.load(wav_path, sr=None)
        duration = librosa.get_duration(y=y, sr=sr)
        
        if duration <= segment_length:
            # Si es corto, transcribir completo
            with sr.AudioFile(wav_path) as source:
                audio_data = recognizer.record(source)
                return recognizer.recognize_google(audio_data, language=language), "Google (archivo completo)"
        
        # Transcribir por segmentos
        segments_text = []
        for i in range(0, int(duration), segment_length):
            start_time = i
            end_time = min(i + segment_length, duration)
            
            # Extraer segmento
            start_sample = int(start_time * sr)
            end_sample = int(end_time * sr)
            segment = y[start_sample:end_sample]
            
            # Crear archivo temporal para el segmento
            temp_segment = tempfile.NamedTemporaryFile(suffix='.wav', delete=False)
            temp_segment.close()
            sf.write(temp_segment.name, segment, sr)
            
            try:
                with sr.AudioFile(temp_segment.name) as source:
                    recognizer.adjust_for_ambient_noise(source, duration=0.5)
                    audio_data = recognizer.record(source)
                    
                try:
                    segment_text = recognizer.recognize_google(audio_data, language=language)
                    segments_text.append(f"[{int(start_time//60)}:{int(start_time%60):02d}-{int(end_time//60)}:{int(end_time%60):02d}] {segment_text}")
                except sr.UnknownValueError:
                    segments_text.append(f"[{int(start_time//60)}:{int(start_time%60):02d}-{int(end_time//60)}:{int(end_time%60):02d}] [Inaudible]")
                except sr.RequestError:
                    segments_text.append(f"[{int(start_time//60)}:{int(start_time%60):02d}-{int(end_time//60)}:{int(end_time%60):02d}] [Error de API]")
                    
            finally:
                try:
                    os.unlink(temp_segment.name)
                except:
                    pass
        
        return "\n".join(segments_text), "Google (por segmentos)"
        
    except Exception as e:
        raise Exception(f"Error en transcripción por segmentos: {e}")

# Función para transcribir múltiples archivos y guardar en TXT
def transcribe_multiple_files(audio_files, output_filename=None, language='es-ES'):
    """Transcribe múltiples archivos de audio y guarda el resultado en un archivo TXT"""
    if output_filename is None:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f'Transcripcion_{timestamp}.txt'
    
    transcriptions = []
    
    with open(output_filename, 'w', encoding='utf-8') as f:
        f.write("TRANSCRIPCIÓN DE ARCHIVOS DE AUDIO\n")
        f.write("=" * 50 + "\n")
        f.write(f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
        f.write(f"Idioma de transcripción: {language}\n")
        f.write(f"Total de archivos: {len(audio_files)}\n\n")
        
        for i, audio_path in enumerate(audio_files, 1):
            file_name = os.path.basename(audio_path)
            f.write(f"\n{'='*60}\n")
            f.write(f"ARCHIVO {i}: {file_name}\n")
            f.write(f"{'='*60}\n")
            
            print(f"Transcribiendo archivo {i}/{len(audio_files)}: {file_name}")
            
            # Obtener duración del archivo
            try:
                y, sr = librosa.load(audio_path, sr=None)
                duration = librosa.get_duration(y=y, sr=sr)
                duration_str = f"{int(duration // 60)}:{int(duration % 60):02d}"
            except:
                duration_str = "Desconocida"
            
            f.write(f"Duración: {duration_str}\n")
            f.write(f"Ruta: {audio_path}\n\n")
            
            # Transcribir el archivo con el idioma especificado
            transcription, method = transcribe_audio(audio_path, language=language)
            
            f.write(f"Método de transcripción: {method}\n")
            f.write(f"TRANSCRIPCIÓN:\n")
            f.write("-" * 40 + "\n")
            f.write(f"{transcription}\n")
            f.write("-" * 40 + "\n")
            
            # Agregar consejos si hay error
            if transcription and transcription.startswith('[Error'):
                f.write("\nCONSEJOS PARA RESOLVER ERRORES:\n")
                f.write("• Verifique que el archivo de audio no esté dañado\n")
                f.write("• Asegúrese de tener conexión estable a internet\n")
                f.write("• Para archivos muy largos, el sistema los segmenta automáticamente\n")
                f.write("• Pruebe con un archivo más corto o de mejor calidad\n\n")
            
            transcriptions.append({
                'file': file_name,
                'path': audio_path,
                'duration': duration_str,
                'transcription': transcription,
                'method': method
            })
    
    print(f"Transcripción completada. Archivo guardado como: {output_filename}")
    return output_filename, transcriptions

# Función para ejecutar el modelo YOLO y retornar detecciones por archivo
def run_yolo_analysis(model_path):
    model = YOLO(model_path)
    input_dir = 'spectrograms'
    output_dir = 'resultados_normal'
    os.makedirs(output_dir, exist_ok=True)
    
    # Diccionario para almacenar detecciones por archivo
    detections_by_file = {}

    for img_name in os.listdir(input_dir):
        if not img_name.lower().endswith(('.png', '.jpg', '.jpeg')):
            continue
            
        img_path = os.path.join(input_dir, img_name)
        
        try:
            img = cv2.imread(img_path)
            if img is None:
                print(f"No se pudo cargar la imagen: {img_name}")
                continue
                
            results = model(img)

            # Extraer información del archivo (prefijo del archivo de audio)
            parts = img_name.split('_')
            
            # Buscar la palabra "spectrogram" para identificar el formato correcto
            spectrogram_index = -1
            for i, part in enumerate(parts):
                if part == 'spectrogram':
                    spectrogram_index = i
                    break
            
            if spectrogram_index >= 0 and len(parts) >= spectrogram_index + 3:
                # El prefijo del archivo es todo lo que está antes de "spectrogram"
                file_prefix = '_'.join(parts[:spectrogram_index])
                try:
                    start_time = int(parts[spectrogram_index + 1])
                    end_time_str = parts[spectrogram_index + 2].replace('.png', '').replace('.jpg', '').replace('.jpeg', '')
                    end_time = int(end_time_str)
                except (ValueError, IndexError) as e:
                    print(f"Error al parsear tiempos del archivo {img_name}: {e}")
                    file_prefix = 'unknown'
                    start_time = 0
                    end_time = 0
            else:
                print(f"Formato de archivo no reconocido: {img_name}")
                file_prefix = 'unknown'
                start_time = 0
                end_time = 0

            if file_prefix not in detections_by_file:
                detections_by_file[file_prefix] = {'segments': [], 'detections': []}

            # Guardar información del segmento
            has_detection = len(results[0].boxes) > 0
            detections_by_file[file_prefix]['segments'].append({
                'start': start_time,
                'end': end_time,
                'has_detection': has_detection,
                'image_path': img_path
            })

            if has_detection:
                shutil.copy(img_path, os.path.join(output_dir, img_name))
                detections_by_file[file_prefix]['detections'].append({
                    'start': start_time,
                    'end': end_time,
                    'image_path': img_path
                })
                
        except Exception as e:
            print(f"Error procesando archivo {img_name}: {e}")
            continue

    print("Análisis YOLO completado.")
    return detections_by_file

# Función para crear gráfico resumen de detecciones
def create_summary_chart(detections_by_file, output_path='summary_chart.png'):
    plt.figure(figsize=(15, 8))
    
    file_names = list(detections_by_file.keys())
    y_positions = range(len(file_names))
    
    colors = {'detection': 'red', 'no_detection': 'green'}
    
    for i, file_name in enumerate(file_names):
        segments = detections_by_file[file_name]['segments']
        
        for segment in segments:
            start = segment['start']
            end = segment['end']
            color = colors['detection'] if segment['has_detection'] else colors['no_detection']
            
            plt.barh(i, end - start, left=start, height=0.6, 
                    color=color, alpha=0.7, edgecolor='black', linewidth=0.5)
    
    plt.yticks(y_positions, file_names)
    plt.xlabel('Tiempo (segundos)')
    plt.ylabel('Archivos de Audio')
    plt.title('Resumen de Detecciones por Archivo de Audio')
    
    # Crear leyenda
    from matplotlib.patches import Patch
    legend_elements = [Patch(facecolor='red', alpha=0.7, label='Corte Detectado'),
                      Patch(facecolor='green', alpha=0.7, label='Sin Corte')]
    plt.legend(handles=legend_elements, loc='upper right')
    
    plt.grid(True, alpha=0.3, axis='x')
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return output_path

# Función para agregar imágenes a una tabla en el documento Word
def add_images_to_table(doc, image_paths, time_interval):
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Image'
    hdr_cells[1].text = f'Time Interval: {time_interval}'

    for image_path in image_paths:
        row_cells = table.add_row().cells
        try:
            img = Image.open(image_path)
            img.thumbnail((Inches(2.0 * 96), Inches(1.0 * 96))) 
            img.save('temp_image.png')
            row_cells[0].paragraphs[0].add_run().add_picture('temp_image.png', width=Inches(2.0))
            file_name = os.path.basename(image_path)
            start, end = file_name.replace('.png', '').split('_')[1:3]
            row_cells[1].text = f'{start} - {end} seconds'
        except Exception as e:
            print(f'Error processing image {image_path}: {e}')
            continue

    if os.path.exists('temp_image.png'):
        os.remove('temp_image.png')

# Función para generar el documento Word con gráfico resumen
def generate_consolidated_report(detections_by_file, audio_files):
    doc = Document()
    doc.add_heading('Reporte Consolidado de Análisis de Audio', level=1)
    
    # Información general
    doc.add_heading('Información General', level=2)
    doc.add_paragraph(f'Fecha de análisis: {datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph(f'Número de archivos analizados: {len(audio_files)}')
    
    # Lista de archivos procesados
    doc.add_heading('Archivos Procesados', level=2)
    for i, audio_file in enumerate(audio_files, 1):
        file_name = os.path.basename(audio_file)
        doc.add_paragraph(f'{i}. {file_name}')
    
    # Crear y agregar gráfico resumen
    chart_path = create_summary_chart(detections_by_file)
    doc.add_heading('Gráfico Resumen de Detecciones', level=2)
    doc.add_paragraph('El siguiente gráfico muestra un resumen de las detecciones encontradas en cada archivo de audio:')
    
    try:
        # Redimensionar imagen para el documento
        img = Image.open(chart_path)
        img.thumbnail((Inches(6.0 * 96), Inches(4.0 * 96)))
        img.save('temp_chart.png')
        doc.add_picture('temp_chart.png', width=Inches(6.0))
        
        if os.path.exists('temp_chart.png'):
            os.remove('temp_chart.png')
    except Exception as e:
        doc.add_paragraph(f'Error al insertar gráfico: {e}')
    
    # Resumen estadístico
    doc.add_heading('Resumen Estadístico', level=2)
    total_segments = 0
    total_detections = 0
    
    for file_prefix, data in detections_by_file.items():
        segments_count = len(data['segments'])
        detections_count = len(data['detections'])
        total_segments += segments_count
        total_detections += detections_count
        
        detection_percentage = (detections_count / segments_count * 100) if segments_count > 0 else 0
        doc.add_paragraph(f'• {file_prefix}: {detections_count}/{segments_count} segmentos con detección ({detection_percentage:.1f}%)')
    
    overall_percentage = (total_detections / total_segments * 100) if total_segments > 0 else 0
    doc.add_paragraph(f'\nResumen general: {total_detections}/{total_segments} segmentos con detección ({overall_percentage:.1f}%)')
    
    # Guardar documento
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    report_name = f'Reporte_Consolidado_{timestamp}.docx'
    doc.save(report_name)
    
    # Limpiar archivo temporal del gráfico
    if os.path.exists(chart_path):
        os.remove(chart_path)
    
    return report_name

class SpectrogramApp(QWidget):
    def __init__(self):
        super().__init__()
        self.audio_files = []  # Lista para almacenar múltiples archivos
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Generador de Espectrogramas y Reportes Consolidados')

        layout = QVBoxLayout()

        # Sección de archivos de audio
        self.label = QLabel('Archivos de audio seleccionados:', self)
        layout.addWidget(self.label)

        # Lista para mostrar archivos seleccionados
        self.audio_list = QListWidget(self)
        self.audio_list.setMaximumHeight(150)
        layout.addWidget(self.audio_list)

        # Botones para manejo de archivos
        file_buttons_layout = QHBoxLayout()
        
        self.add_files_button = QPushButton('Agregar archivos', self)
        self.add_files_button.clicked.connect(self.add_audio_files)
        file_buttons_layout.addWidget(self.add_files_button)
        
        self.clear_files_button = QPushButton('Limpiar lista', self)
        self.clear_files_button.clicked.connect(self.clear_audio_files)
        file_buttons_layout.addWidget(self.clear_files_button)
        
        self.remove_file_button = QPushButton('Eliminar seleccionado', self)
        self.remove_file_button.clicked.connect(self.remove_selected_file)
        file_buttons_layout.addWidget(self.remove_file_button)
        
        layout.addLayout(file_buttons_layout)

        # Sección de configuración de duración de segmentos
        segment_config_label = QLabel('Configuración de espectrogramas:', self)
        segment_config_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(segment_config_label)
        
        segment_layout = QHBoxLayout()
        
        self.segment_label = QLabel('Duración de cada espectrograma:', self)
        segment_layout.addWidget(self.segment_label)
        
        self.segment_duration_combo = QComboBox(self)
        self.segment_duration_combo.addItems([
            '1 segundo',
            '3 segundos', 
            '5 segundos',
            '10 segundos',
            '15 segundos',
            '30 segundos',
            '1 minuto (60 segundos)',
            '2 minutos (120 segundos)',
            '3 minutos (180 segundos)',
            '5 minutos (300 segundos)'
        ])
        self.segment_duration_combo.setCurrentText('3 segundos')  # Valor por defecto
        segment_layout.addWidget(self.segment_duration_combo)
        
        layout.addLayout(segment_layout)
        
        # Información sobre duraciones
        duration_info_label = QLabel('• Segmentos más cortos: Mayor precisión en la detección\n'
                                    '• Segmentos más largos: Visión más amplia del contexto\n'
                                    '• Para análisis detallado usar 1-5 segundos\n'
                                    '• Para análisis general usar 1-5 minutos', self)
        duration_info_label.setStyleSheet("font-size: 9px; color: #666; margin: 5px;")
        layout.addWidget(duration_info_label)

        # Sección para generación por rango de tiempo
        range_label = QLabel('Generar por rango de tiempo (se exporta en directorio separado):', self)
        layout.addWidget(range_label)
        
        time_range_layout = QHBoxLayout()
        
        self.start_time_label = QLabel('Tiempo inicial (seg):', self)
        time_range_layout.addWidget(self.start_time_label)
        
        self.start_time_input = QLineEdit(self)
        self.start_time_input.setPlaceholderText("0")
        time_range_layout.addWidget(self.start_time_input)
        
        self.end_time_label = QLabel('Tiempo final (seg):', self)
        time_range_layout.addWidget(self.end_time_label)
        
        self.end_time_input = QLineEdit(self)
        self.end_time_input.setPlaceholderText("30")
        time_range_layout.addWidget(self.end_time_input)
        
        layout.addLayout(time_range_layout)
        
        # Nueva sección para modo de generación de rango de tiempo
        time_mode_label = QLabel('Modo de generación por rango:', self)
        layout.addWidget(time_mode_label)
        
        self.time_mode_group = QButtonGroup(self)
        self.radio_complete_range = QRadioButton('Rango completo (cobertura total del período)', self)
        self.radio_combined_mode = QRadioButton('Combinado (rango completo + saltos adicionales)', self)
        self.radio_complete_range.setChecked(True)  # Por defecto cobertura completa
        layout.addWidget(self.radio_complete_range)
        layout.addWidget(self.radio_combined_mode)
        self.time_mode_group.addButton(self.radio_complete_range)
        self.time_mode_group.addButton(self.radio_combined_mode)
        
        # Conectar señales para habilitar/deshabilitar el campo de salto
        self.radio_complete_range.toggled.connect(self.toggle_time_jump_input)
        self.radio_combined_mode.toggled.connect(self.toggle_time_jump_input)
        
        # Sección para configurar salto de tiempo para modo combinado
        time_jump_layout = QHBoxLayout()
        
        self.time_jump_label = QLabel('Salto adicional (seg):', self)
        time_jump_layout.addWidget(self.time_jump_label)
        
        self.time_jump_input = QLineEdit(self)
        self.time_jump_input.setPlaceholderText("3")
        self.time_jump_input.setText("3")  # Valor por defecto
        self.time_jump_input.setEnabled(False)  # Deshabilitado por defecto
        time_jump_layout.addWidget(self.time_jump_input)
        
        layout.addLayout(time_jump_layout)

        # Sección separada para generación por saltos únicamente
        jumps_section_label = QLabel('Generar espectrogramas por saltos únicamente:', self)
        jumps_section_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(jumps_section_label)
        
        jump_config_layout = QHBoxLayout()
        
        self.jump_interval_label = QLabel('Intervalo de salto (seg):', self)
        jump_config_layout.addWidget(self.jump_interval_label)
        
        self.jump_interval_input = QLineEdit(self)
        self.jump_interval_input.setPlaceholderText("5")
        self.jump_interval_input.setText("5")  # Valor por defecto diferente
        jump_config_layout.addWidget(self.jump_interval_input)
        
        layout.addLayout(jump_config_layout)
        
        # Agregar información sobre los modos
        mode_info_label = QLabel('• Rango completo: Analiza todo el período especificado sin saltos\n'
                                '• Combinado: Análisis completo del rango + saltos adicionales\n'
                                '• Por saltos: Analiza todo el audio solo en intervalos específicos', self)
        mode_info_label.setStyleSheet("font-size: 9px; color: #666; margin: 5px;")
        layout.addWidget(mode_info_label)

        # Botones de procesamiento
        self.generate_button = QPushButton('Generar Espectrogramas para todos los archivos', self)
        self.generate_button.clicked.connect(self.generate_all_spectrograms)
        layout.addWidget(self.generate_button)
        
        self.generate_range_button = QPushButton('Generar Espectrogramas por Rango de Tiempo', self)
        self.generate_range_button.clicked.connect(self.generate_spectrograms_by_range)
        layout.addWidget(self.generate_range_button)
        
        self.generate_jumps_button = QPushButton('Generar Espectrogramas por Saltos (directorio separado)', self)
        self.generate_jumps_button.clicked.connect(self.generate_spectrograms_by_jumps)
        layout.addWidget(self.generate_jumps_button)

        self.transcribe_button = QPushButton('Transcribir Archivos de Audio', self)
        self.transcribe_button.clicked.connect(self.transcribe_audio_files)
        layout.addWidget(self.transcribe_button)
        
        # Configuración de transcripción
        transcription_config_layout = QHBoxLayout()
        
        self.language_label = QLabel('Idioma para transcripción:', self)
        transcription_config_layout.addWidget(self.language_label)
        
        self.language_combo = QComboBox(self)
        self.language_combo.addItems([
            'es-ES (Español España)', 
            'es-MX (Español México)', 
            'en-US (Inglés)', 
            'fr-FR (Francés)',
            'it-IT (Italiano)',
            'pt-PT (Portugués)',
            'de-DE (Alemán)'
        ])
        transcription_config_layout.addWidget(self.language_combo)
        
        layout.addLayout(transcription_config_layout)

        self.yolo_button = QPushButton('Ejecutar análisis de autentificación', self)
        self.yolo_button.clicked.connect(self.run_yolo_analysis)
        layout.addWidget(self.yolo_button)

        self.report_button = QPushButton('Generar Reporte Consolidado', self)
        self.report_button.clicked.connect(self.generate_consolidated_report)
        layout.addWidget(self.report_button)

        self.setLayout(layout)

    def add_audio_files(self):
        options = QFileDialog.Options()
        file_paths, _ = QFileDialog.getOpenFileNames(
            self, 
            "Seleccionar archivos de audio", 
            "", 
            "Audio Files (*.wav *.mp3 *.flac *.m4a);;All Files (*)", 
            options=options
        )
        
        for file_path in file_paths:
            if file_path not in self.audio_files:
                self.audio_files.append(file_path)
                self.audio_list.addItem(os.path.basename(file_path))

    def clear_audio_files(self):
        self.audio_files.clear()
        self.audio_list.clear()

    def remove_selected_file(self):
        current_row = self.audio_list.currentRow()
        if current_row >= 0:
            self.audio_files.pop(current_row)
            self.audio_list.takeItem(current_row)

    def get_selected_segment_duration(self):
        """Obtener la duración del segmento seleccionada en segundos"""
        duration_text = self.segment_duration_combo.currentText()
        
        # Extraer el número de la cadena
        if '1 segundo' in duration_text:
            return 1
        elif '3 segundos' in duration_text:
            return 3
        elif '5 segundos' in duration_text:
            return 5
        elif '10 segundos' in duration_text:
            return 10
        elif '15 segundos' in duration_text:
            return 15
        elif '30 segundos' in duration_text:
            return 30
        elif '1 minuto' in duration_text:
            return 60
        elif '2 minutos' in duration_text:
            return 120
        elif '3 minutos' in duration_text:
            return 180
        elif '5 minutos' in duration_text:
            return 300
        else:
            return 3  # Valor por defecto

    def generate_all_spectrograms(self):
        if not self.audio_files:
            QMessageBox.warning(self, "Advertencia", "Por favor, selecciona al menos un archivo de audio.")
            return

        segment_length = self.get_selected_segment_duration()
        
        # Limpiar directorio de espectrogramas antes de empezar
        if os.path.exists('spectrograms'):
            shutil.rmtree('spectrograms')
        
        try:
            for i, audio_path in enumerate(self.audio_files):
                file_name = os.path.basename(audio_path)
                file_prefix = f"audio_{i+1}_{os.path.splitext(file_name)[0]}"
                generate_spectrograms_for_file(audio_path, file_prefix, segment_length)
            
            duration_text = self.segment_duration_combo.currentText()
            QMessageBox.information(self, "Éxito", 
                                  f"Espectrogramas generados para {len(self.audio_files)} archivos.\n"
                                  f"Duración de cada espectrograma: {duration_text}\n"
                                  f"Directorio: spectrograms/")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al generar espectrogramas: {str(e)}")

    def toggle_time_jump_input(self):
        """Habilitar/deshabilitar el campo de salto de tiempo según el modo seleccionado"""
        needs_jump_input = self.radio_combined_mode.isChecked()
        self.time_jump_input.setEnabled(needs_jump_input)
        self.time_jump_label.setEnabled(needs_jump_input)

    def generate_spectrograms_by_range(self):
        if not self.audio_files:
            QMessageBox.warning(self, "Advertencia", "Por favor, selecciona al menos un archivo de audio.")
            return

        # Validar y obtener los tiempos
        try:
            start_time_str = self.start_time_input.text().strip()
            end_time_str = self.end_time_input.text().strip()
            
            if not start_time_str or not end_time_str:
                QMessageBox.warning(self, "Advertencia", "Por favor, ingresa tanto el tiempo inicial como el final.")
                return
            
            start_time = float(start_time_str)
            end_time = float(end_time_str)
            
            # Determinar el modo y configuración según las selecciones
            if self.radio_complete_range.isChecked():
                mode = "complete"
                time_jump = None
                mode_display = "Rango completo (cobertura total)"
            else:  # combined mode
                mode = "combined"
                time_jump_str = self.time_jump_input.text().strip()
                if not time_jump_str:
                    time_jump = 3.0
                else:
                    time_jump = float(time_jump_str)
                    if time_jump <= 0:
                        QMessageBox.warning(self, "Advertencia", "El salto de tiempo debe ser mayor que 0.")
                        return
                mode_display = f"Combinado (completo + saltos cada {time_jump}s)"
            
            if start_time < 0:
                QMessageBox.warning(self, "Advertencia", "El tiempo inicial debe ser mayor o igual a 0.")
                return
            
            if start_time >= end_time:
                QMessageBox.warning(self, "Advertencia", "El tiempo inicial debe ser menor que el tiempo final.")
                return
                
        except ValueError:
            QMessageBox.warning(self, "Advertencia", "Por favor, ingresa valores numéricos válidos para los tiempos.")
            return

        segment_length = self.get_selected_segment_duration()
        
        # Limpiar directorio de espectrogramas de rango de tiempo antes de empezar
        if os.path.exists('spectrograms_time_range'):
            shutil.rmtree('spectrograms_time_range')
        
        try:
            for i, audio_path in enumerate(self.audio_files):
                file_name = os.path.basename(audio_path)
                file_prefix = f"audio_{i+1}_{os.path.splitext(file_name)[0]}"
                
                # Verificar que el archivo tenga suficiente duración
                y, sr = librosa.load(audio_path, sr=None)
                duration = librosa.get_duration(y=y, sr=sr)
                
                if start_time >= duration:
                    QMessageBox.warning(self, "Advertencia", 
                                      f"El archivo {file_name} tiene una duración de {duration:.2f}s, "
                                      f"que es menor al tiempo inicial especificado ({start_time}s).")
                    continue
                
                actual_end_time = min(end_time, duration)
                generate_spectrograms_by_time_range(audio_path, file_prefix, start_time, actual_end_time, 
                                                  segment_length, mode, time_jump if time_jump else 3)
            
            # Calcular número estimado de espectrogramas por archivo
            if mode == "complete":
                estimated_spectrograms = int((end_time - start_time) / segment_length) + 1
            else:  # combined
                complete_specs = int((end_time - start_time) / segment_length) + 1
                jump_specs = int((end_time - start_time) / time_jump) + 1
                estimated_spectrograms = complete_specs + jump_specs
            
            duration_text = self.segment_duration_combo.currentText()
            QMessageBox.information(self, "Éxito", 
                                  f"Espectrogramas generados para {len(self.audio_files)} archivos "
                                  f"en el rango {start_time}s - {end_time}s.\n"
                                  f"Modo: {mode_display}\n"
                                  f"Duración de cada espectrograma: {duration_text}\n"
                                  f"Directorio: spectrograms_time_range/\n"
                                  f"Espectrogramas estimados por archivo: ~{estimated_spectrograms}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al generar espectrogramas por rango: {str(e)}")

    def generate_spectrograms_by_jumps(self):
        if not self.audio_files:
            QMessageBox.warning(self, "Advertencia", "Por favor, selecciona al menos un archivo de audio.")
            return

        # Validar y obtener el intervalo de salto
        try:
            jump_interval_str = self.jump_interval_input.text().strip()
            
            if not jump_interval_str:
                jump_interval = 5.0  # Valor por defecto
            else:
                jump_interval = float(jump_interval_str)
                if jump_interval <= 0:
                    QMessageBox.warning(self, "Advertencia", "El intervalo de salto debe ser mayor que 0.")
                    return
                    
        except ValueError:
            QMessageBox.warning(self, "Advertencia", "Por favor, ingresa un valor numérico válido para el intervalo de salto.")
            return

        segment_length = self.get_selected_segment_duration()
        
        # Limpiar directorio de espectrogramas por saltos antes de empezar
        if os.path.exists('spectrograms_jumps'):
            shutil.rmtree('spectrograms_jumps')
        
        try:
            for i, audio_path in enumerate(self.audio_files):
                file_name = os.path.basename(audio_path)
                file_prefix = f"audio_{i+1}_{os.path.splitext(file_name)[0]}"
                
                # Obtener duración del archivo para estimación
                y, sr = librosa.load(audio_path, sr=None)
                duration = librosa.get_duration(y=y, sr=sr)
                
                generate_spectrograms_by_jumps(audio_path, file_prefix, jump_interval, segment_length)
            
            # Calcular número estimado de espectrogramas por archivo
            estimated_spectrograms = int(duration / jump_interval) + 1
            
            duration_text = self.segment_duration_combo.currentText()
            QMessageBox.information(self, "Éxito", 
                                  f"Espectrogramas por saltos generados para {len(self.audio_files)} archivos.\n"
                                  f"Intervalo de salto: {jump_interval}s\n"
                                  f"Duración de cada espectrograma: {duration_text}\n"
                                  f"Directorio: spectrograms_jumps/\n"
                                  f"Espectrogramas estimados por archivo: ~{estimated_spectrograms}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al generar espectrogramas por saltos: {str(e)}")

    def transcribe_audio_files(self):
        if not self.audio_files:
            QMessageBox.warning(self, "Advertencia", "Por favor, selecciona al menos un archivo de audio.")
            return
        
        # Obtener idioma seleccionado
        language_text = self.language_combo.currentText()
        language_code = language_text.split(' ')[0]  # Extraer el código (ej: 'es-ES')
        
        # Confirmar la operación ya que puede tomar tiempo
        reply = QMessageBox.question(self, 'Confirmar Transcripción', 
                                   f'¿Deseas transcribir {len(self.audio_files)} archivo(s) de audio?\n\n'
                                   f'Idioma: {language_text}\n'
                                   'Nota: Este proceso puede tomar varios minutos y requiere conexión a internet '
                                   'para obtener los mejores resultados.\n\n'
                                   'El sistema intentará múltiples métodos si hay errores.',
                                   QMessageBox.Yes | QMessageBox.No, 
                                   QMessageBox.No)
        
        if reply == QMessageBox.No:
            return
        
        try:
            # Mostrar mensaje de progreso
            progress_msg = QMessageBox(self)
            progress_msg.setWindowTitle("Transcribiendo...")
            progress_msg.setText("Transcribiendo archivos de audio...\nEsto puede tomar varios minutos.\n\nUsando múltiples métodos para mejor precisión.")
            progress_msg.setStandardButtons(QMessageBox.NoButton)
            progress_msg.show()
            
            # Procesar la aplicación para mostrar el mensaje
            QApplication.processEvents()
            
            # Realizar la transcripción con el idioma seleccionado
            output_filename, transcriptions = transcribe_multiple_files(self.audio_files, language=language_code)
            
            # Cerrar mensaje de progreso
            progress_msg.close()
            
            # Mostrar resumen de resultados
            successful_transcriptions = sum(1 for t in transcriptions 
                                          if not t['transcription'].startswith('['))
            
            QMessageBox.information(self, "Transcripción Completada", 
                                  f"Transcripción completada!\n\n"
                                  f"Archivos procesados: {len(self.audio_files)}\n"
                                  f"Transcripciones exitosas: {successful_transcriptions}\n"
                                  f"Idioma usado: {language_text}\n"
                                  f"Archivo guardado como: {output_filename}")
                                  
        except Exception as e:
            progress_msg.close()
            QMessageBox.critical(self, "Error", f"Error durante la transcripción: {str(e)}")
            
    # Nota: Si tienes el error 'Bad Request', puede deberse a:
    # 1. Archivo de audio dañado o en formato incompatible
    # 2. Archivo demasiado largo para Google API
    # 3. Problemas de conexión a internet  
    # 4. El archivo tiene características que Google no puede procesar

    def run_yolo_analysis(self):
        if not os.path.exists('spectrograms') or not os.listdir('spectrograms'):
            QMessageBox.warning(self, "Advertencia", "Primero genera los espectrogramas.")
            return
        
        try:
            # Seleccionar modelo basado en la duración del segmento
            segment_length = self.get_selected_segment_duration()
            if segment_length == 1:
                model_path = 'modelos/grayscale/best.pt'
            else:
                model_path = 'modelos/normal/best.pt'
                
            self.detections_data = run_yolo_analysis(model_path)
            QMessageBox.information(self, "Análisis completo", "El análisis de autentificación ha finalizado.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error en el análisis YOLO: {str(e)}")

    def generate_consolidated_report(self):
        if not hasattr(self, 'detections_data') or not self.detections_data:
            QMessageBox.warning(self, "Advertencia", "Primero ejecuta el análisis de autentificación.")
            return
        
        try:
            report_name = generate_consolidated_report(self.detections_data, self.audio_files)
            QMessageBox.information(self, "Éxito", f"Reporte consolidado generado: {report_name}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al generar reporte: {str(e)}")

def main():
    app = QApplication(sys.argv)
    ex = SpectrogramApp()
    ex.show()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()
